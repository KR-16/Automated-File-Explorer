import os
from fpdf import FPDF
from docx import Document

# Define project structure
project_structure = {
    "root": "docs",
    "files": [
        "faq.txt",
        "config_guide.pdf",
        "api_docs.docx"
    ]
}

# Create root directory
os.makedirs(project_structure["root"], exist_ok=True)

# Create faq.txt
faq_content = """Frequently Asked Questions (FAQ)

1. What is this project about?
   - This project provides documentation for API usage and configuration.

2. How do I install the necessary dependencies?
   - Follow the steps in config_guide.pdf.

3. Where can I find API details?
   - Refer to api_docs.docx.

For further queries, contact support@example.com.
"""
faq_path = os.path.join(project_structure["root"], "faq.txt")
with open(faq_path, "w") as f:
    f.write(faq_content)

# Create config_guide.pdf
pdf = FPDF()
pdf.set_auto_page_break(auto=True, margin=15)
pdf.add_page()
pdf.set_font("Arial", size=12)
pdf.cell(200, 10, "Configuration Guide", ln=True, align="C")
pdf.ln(10)
pdf.multi_cell(0, 10, "This document provides guidance on setting up and configuring the system.\n\n"
                      "Steps:\n"
                      "1. Install dependencies.\n"
                      "2. Set up environment variables.\n"
                      "3. Run the initialization script.\n"
                      "4. Test the setup to ensure proper functionality.\n\n"
                      "For detailed instructions, refer to api_docs.docx.")
config_guide_path = os.path.join(project_structure["root"], "config_guide.pdf")
pdf.output(config_guide_path)

# Create api_docs.docx
doc = Document()
doc.add_heading("API Documentation", level=1)

doc.add_heading("Introduction", level=2)
doc.add_paragraph("This document provides details about the API endpoints and their usage.")

doc.add_heading("Endpoints", level=2)
doc.add_paragraph("1. GET /api/status - Checks the status of the service.\n"
                  "2. POST /api/data - Submits data to the server.\n"
                  "3. DELETE /api/data/{id} - Deletes a specific data entry.")

doc.add_heading("Authentication", level=2)
doc.add_paragraph("API requests require an authentication token. Use the following format:\n"
                  "Authorization: Bearer <token>")

api_docs_path = os.path.join(project_structure["root"], "api_docs.docx")
doc.save(api_docs_path)

print("Project structure created successfully!")
